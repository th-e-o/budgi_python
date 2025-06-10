# modules/pdf_to_word_converter.py
import os
from pathlib import Path
from typing import Optional, Union, Dict
import logging
from datetime import datetime
import tempfile

# Import des bibliothèques nécessaires
from pdf2docx import Converter
import PyPDF2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class PDFToWordConverter:
    """Convertit des fichiers PDF en documents Word"""
    
    def __init__(self):
        self.temp_files = []
        
    def convert_pdf_to_docx(self, pdf_path: Union[str, Path], 
                           output_path: Optional[Union[str, Path]] = None,
                           preserve_layout: bool = True) -> Optional[str]:
        """
        Convertit un fichier PDF en DOCX
        
        Args:
            pdf_path: Chemin vers le fichier PDF
            output_path: Chemin de sortie (optionnel)
            preserve_layout: Préserver la mise en forme (True) ou extraire le texte brut (False)
            
        Returns:
            Chemin vers le fichier DOCX créé ou None en cas d'erreur
        """
        try:
            pdf_path = Path(pdf_path)
            
            if not pdf_path.exists():
                logger.error(f"Fichier PDF non trouvé: {pdf_path}")
                return None
                
            # Définir le chemin de sortie
            if output_path is None:
                # Créer un fichier temporaire
                fd, temp_path = tempfile.mkstemp(suffix='.docx')
                os.close(fd)
                output_path = Path(temp_path)
                self.temp_files.append(temp_path)
            else:
                output_path = Path(output_path)
            
            if preserve_layout:
                # Utiliser pdf2docx pour préserver la mise en forme
                return self._convert_with_layout(pdf_path, output_path)
            else:
                # Extraire le texte brut et créer un document simple
                return self._convert_text_only(pdf_path, output_path)
                
        except Exception as e:
            logger.error(f"Erreur conversion PDF->DOCX: {str(e)}")
            return None
    
    def _convert_with_layout(self, pdf_path: Path, output_path: Path) -> Optional[str]:
        """Convertit en préservant la mise en forme avec pdf2docx"""
        try:
            # Créer le convertisseur
            cv = Converter(str(pdf_path))
            
            # Convertir toutes les pages
            cv.convert(str(output_path), start=0, end=None)
            
            # Fermer le convertisseur
            cv.close()
            
            logger.info(f"PDF converti avec mise en forme: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Erreur conversion avec layout: {str(e)}")
            # Fallback sur la conversion texte simple
            return self._convert_text_only(pdf_path, output_path)
    
    def _convert_text_only(self, pdf_path: Path, output_path: Path) -> Optional[str]:
        """Convertit en extrayant uniquement le texte"""
        try:
            # Extraire le texte du PDF
            text_content = self._extract_pdf_text(pdf_path)
            
            if not text_content:
                logger.warning("Aucun texte extrait du PDF")
                return None
            
            # Créer un document Word
            doc = Document()
            
            # Ajouter un titre avec le nom du fichier
            title = doc.add_heading(f'Document converti depuis: {pdf_path.name}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajouter la date de conversion
            date_para = doc.add_paragraph(f'Date de conversion: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajouter une ligne de séparation
            doc.add_paragraph('_' * 50)
            
            # Traiter le texte par pages
            pages = text_content.split('\n\n[Page ')
            
            for i, page_content in enumerate(pages):
                if i > 0:
                    # Ajouter un saut de page entre les pages
                    doc.add_page_break()
                    # Ajouter le numéro de page
                    doc.add_heading(f'Page {i + 1}', level=2)
                
                # Traiter le contenu de la page
                paragraphs = page_content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Nettoyer le texte
                        clean_text = para_text.strip()
                        
                        # Détecter les titres potentiels (lignes courtes en majuscules)
                        if len(clean_text) < 100 and clean_text.isupper():
                            doc.add_heading(clean_text, level=1)
                        else:
                            # Ajouter comme paragraphe normal
                            para = doc.add_paragraph(clean_text)
                            # Justifier le texte
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Sauvegarder le document
            doc.save(str(output_path))
            
            logger.info(f"PDF converti (texte seul): {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Erreur conversion texte: {str(e)}")
            return None
    
    def _extract_pdf_text(self, pdf_path: Path) -> str:
        """Extrait le texte d'un PDF avec PyPDF2"""
        try:
            text_content = []
            
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Extraire le texte de chaque page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            if page_num > 0:
                                text_content.append(f"\n\n[Page {page_num + 1}]\n")
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Erreur extraction page {page_num + 1}: {str(e)}")
                        continue
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Erreur extraction texte PDF: {str(e)}")
            return ""
    
    def convert_pdf_bytes_to_docx(self, pdf_bytes: bytes, 
                                preserve_layout: bool = True) -> Optional[bytes]:
        """
        Convertit des bytes PDF en bytes DOCX
        
        Args:
            pdf_bytes: Contenu du PDF en bytes
            preserve_layout: Préserver la mise en forme
            
        Returns:
            Contenu DOCX en bytes ou None
        """
        try:
            # Créer des fichiers temporaires
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_tmp:
                pdf_tmp.write(pdf_bytes)
                pdf_path = pdf_tmp.name
            
            self.temp_files.append(pdf_path)
            
            # Convertir
            docx_path = self.convert_pdf_to_docx(pdf_path, preserve_layout=preserve_layout)
            
            if docx_path:
                # Lire le fichier DOCX
                with open(docx_path, 'rb') as docx_file:
                    docx_bytes = docx_file.read()
                return docx_bytes
            
            return None
            
        except Exception as e:
            logger.error(f"Erreur conversion bytes: {str(e)}")
            return None
        finally:
            # Nettoyer le fichier PDF temporaire
            if 'pdf_path' in locals() and os.path.exists(pdf_path):
                try:
                    os.unlink(pdf_path)
                except:
                    pass
    
    def get_pdf_info(self, pdf_path: Union[str, Path]) -> Dict[str, any]:
        """Récupère des informations sur le PDF"""
        try:
            pdf_path = Path(pdf_path)
            info = {
                'filename': pdf_path.name,
                'size_mb': pdf_path.stat().st_size / (1024 * 1024),
                'pages': 0,
                'has_text': False,
                'has_images': False
            }
            
            with open(pdf_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                info['pages'] = len(pdf_reader.pages)
                
                # Vérifier s'il y a du texte
                for page in pdf_reader.pages[:5]:  # Vérifier les 5 premières pages
                    if page.extract_text().strip():
                        info['has_text'] = True
                        break
                
                # Vérifier les métadonnées
                if pdf_reader.metadata:
                    info['metadata'] = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    }
            
            return info
            
        except Exception as e:
            logger.error(f"Erreur lecture info PDF: {str(e)}")
            return {'error': str(e)}
    
    def cleanup_temp_files(self):
        """Nettoie les fichiers temporaires"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
                    logger.debug(f"Fichier temporaire supprimé: {temp_file}")
            except Exception as e:
                logger.warning(f"Impossible de supprimer {temp_file}: {str(e)}")
        self.temp_files = []
    
    def __del__(self):
        """Destructeur pour nettoyer les fichiers temporaires"""
        self.cleanup_temp_files()